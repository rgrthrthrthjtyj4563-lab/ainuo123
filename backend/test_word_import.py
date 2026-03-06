from docx import Document
from io import BytesIO
from app.services.word_extractor import WordStructureExtractor

def create_test_docx():
    doc = Document()
    doc.add_heading('文档标题', 0)
    doc.add_heading('一级标题 (Heading 1)', 1)
    doc.add_paragraph('正文内容...')
    doc.add_heading('二级标题 (Heading 2)', 2)
    doc.add_paragraph('正文内容...')
    doc.add_heading('三级标题 (Heading 3)', 3)
    
    # Add some custom style paragraphs if possible, or just standard ones
    p = doc.add_paragraph('模拟标题 1 样式')
    p.style = 'Heading 1'
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_fallback_test_docx():
    doc = Document()
    p1 = doc.add_paragraph('加粗的短文本作为标题')
    p1.runs[0].bold = True
    
    doc.add_paragraph('这是正文内容，不加粗。')
    
    p2 = doc.add_paragraph('另一个加粗标题')
    p2.runs[0].bold = True
    
    doc.add_paragraph('这是正文内容。')
    
    # Long bold paragraph (should NOT be a heading)
    long_text = '这是一段非常非常长的加粗文本，它不应该被识别为标题，因为它太长了。' * 5
    p3 = doc.add_paragraph(long_text)
    if p3.runs:
        p3.runs[0].bold = True
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def test_extraction():
    print("--- Test Standard Styles ---")
    content = create_test_docx()
    extractor = WordStructureExtractor()
    try:
        result = extractor.extract(content, "test.docx")
        print_structure(result['structure'])
    except Exception as e:
        print(f"Error: {e}")

    print("\n--- Test Fallback Strategy ---")
    content_fallback = create_fallback_test_docx()
    extractor_fallback = WordStructureExtractor()
    try:
        result = extractor_fallback.extract(content_fallback, "fallback.docx")
        print_structure(result['structure'])
    except Exception as e:
        print(f"Error: {e}")

def print_structure(nodes, indent=0):
    if not nodes:
        return
    for node in nodes:
        print(f"{'  ' * indent}- {node['title']} (Level {node['level']}, Type: {node['node_type']})")
        if node.get('children'):
            print_structure(node['children'], indent + 1)

if __name__ == "__main__":
    test_extraction()
