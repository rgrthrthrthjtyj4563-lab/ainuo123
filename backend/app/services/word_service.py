"""
Word报告生成服务
"""
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from typing import Dict, Any, List, Optional
from config import get_settings

settings = get_settings()


class WordService:
    """Word文档生成服务"""
    
    def generate_report(self, report_data: Dict[str, Any], output_path: str) -> str:
        """
        生成Word报告
        """
        doc = Document()
        
        # 设置默认样式
        self._set_default_styles(doc)
        
        # 1. 前言
        self._add_section(doc, '前言', report_data.get('generated_content', {}).get('preface', ''))
        
        # 2. 项目背景
        self._add_section(doc, '一、项目背景', report_data.get('generated_content', {}).get('background', ''))
        
        # 3. 项目开展情况
        self._add_section(doc, '二、项目开展情况', report_data.get('generated_content', {}).get('project_info', ''))
        
        # 4. 问卷说明（固定模板）
        self._add_survey_description(doc)
        
        # 5. 问卷结果分析
        self._add_dimension_analysis(doc, report_data)
        
        # 6. 调研结果
        self._add_section(doc, '五、调研结果', report_data.get('generated_content', {}).get('conclusion', ''))
        
        # 7. 建议
        self._add_section(doc, '六、建议', report_data.get('generated_content', {}).get('suggestions', ''))
        
        # 8. 附件
        self._add_attachments(doc, report_data)
        
        # 保存文档
        doc.save(output_path)
        
        return output_path
    
    def _set_default_styles(self, doc: Document):
        """设置默认样式"""
        # 设置Normal样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _add_section(self, doc: Document, title: str, content: str):
        """添加章节"""
        # 添加标题
        heading = doc.add_heading(title, level=1)
        self._set_heading_style(heading, Pt(16))
        
        # 添加内容
        if content:
            for paragraph_text in content.split('\n\n'):
                if paragraph_text.strip():
                    p = doc.add_paragraph(paragraph_text.strip())
                    self._set_paragraph_style(p)
    
    def _add_survey_description(self, doc: Document):
        """添加问卷说明（固定模板）"""
        heading = doc.add_heading('三、问卷说明', level=1)
        self._set_heading_style(heading, Pt(16))
        
        content = """为增强数据间的可比性与分析结果的准确性，本次报告对问卷原始数据采用以下标准化统计处理方法进行系统整理：

1. 数据清洗：严格筛选原始数据，剔除无效填写、关键信息缺失等不符合调研规范的数据样本，保障分析数据的真实性与可靠性；

2. 数据转换：针对问卷收集的分类数据进行规范化转换，使其适配后续统计分析逻辑，提升数据应用的合理性与一致性；

3. 数据聚合：对具有同类属性的调研数据进行归类整合，凝练核心分析维度，避免数据分散导致的结论偏差，聚焦关键研究问题；

4. 数据可视化：结合报告呈现需求，通过直观的可视化方式展示数据分布特征与占比关系，助力快速把握数据规律与核心信息。

通过上述标准化数据处理流程，有效提升了数据的可比性与分析的严谨性，为报告结论的科学性与可信度提供了坚实支撑。"""
        
        for paragraph_text in content.split('\n\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                self._set_paragraph_style(p)

    def _generate_chart(self, question_data: Dict[str, Any], chart_config: Dict[str, Any]) -> Optional[str]:
        """
        生成图表图片，返回临时文件路径
        """
        try:
            import matplotlib.pyplot as plt
            import matplotlib
            
            # 设置中文字体
            # 尝试使用系统常见中文字体
            fonts = ['SimHei', 'Arial Unicode MS', 'PingFang SC', 'Microsoft YaHei', 'Heiti TC']
            for font in fonts:
                try:
                    matplotlib.rcParams['font.sans-serif'] = [font]
                    # 简单测试一下是否可用
                    # plt.figure()
                    # plt.close()
                    break
                except:
                    continue
            
            matplotlib.rcParams['axes.unicode_minus'] = False
            
            options = question_data.get('options', [])
            if not options:
                return None
            
            chart_type = chart_config.get('chart_type', 'bar')
            
            # 准备数据
            # 截短标签，避免过长
            labels = [opt['text'][:15] + '...' if len(opt['text']) > 15 else opt['text'] for opt in options]
            
            # 处理百分比数据，去除%号
            values = []
            for opt in options:
                val_str = str(opt.get('percentage', '0')).rstrip('%')
                try:
                    values.append(float(val_str))
                except ValueError:
                    values.append(0)
            
            # 创建图表
            plt.figure(figsize=(8, 5))
            fig, ax = plt.subplots(figsize=(8, 5))
            
            if chart_type == 'pie':
                # 饼图
                colors = plt.cm.Set3(range(len(labels)))
                wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
                ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
                # 设置字体大小
                plt.setp(texts, size=9)
                plt.setp(autotexts, size=9)
                
            elif chart_type == 'bar':
                # 柱状图
                bars = ax.bar(labels, values, color='#6366F1')
                ax.set_ylabel('占比 (%)', fontsize=10)
                # ax.set_title(question_data.get('question_title', ''), fontsize=12, pad=20)
                plt.xticks(rotation=15, ha='right')
                # 在柱子上显示数值
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{height:.1f}%', ha='center', va='bottom', fontsize=9)
                           
            elif chart_type == 'horizontal_bar':
                # 条形图
                y_pos = range(len(labels))
                ax.barh(y_pos, values, color='#6366F1')
                ax.set_yticks(y_pos)
                ax.set_yticklabels(labels)
                ax.invert_yaxis()  # labels read top-to-bottom
                ax.set_xlabel('占比 (%)', fontsize=10)
                # 在条形旁显示数值
                for i, v in enumerate(values):
                    ax.text(v + 1, i, f'{v:.1f}%', va='center', fontsize=9)
            
            elif chart_type == 'radar':
                # 雷达图 (简单实现)
                import numpy as np
                # Number of variables
                N = len(labels)
                if N > 2: # Radar chart needs at least 3 variables
                    # What will be the angle of each axis in the plot? (we divide the plot / number of variable)
                    angles = [n / float(N) * 2 * np.pi for n in range(N)]
                    angles += angles[:1]
                    
                    values_radar = values + values[:1]
                    
                    ax = plt.subplot(111, polar=True)
                    
                    # Draw one axe per variable + add labels
                    plt.xticks(angles[:-1], labels, color='grey', size=8)
                    
                    # Draw ylabels
                    ax.set_rlabel_position(0)
                    
                    # Plot data
                    ax.plot(angles, values_radar, linewidth=1, linestyle='solid')
                    
                    # Fill area
                    ax.fill(angles, values_radar, 'b', alpha=0.1)

            # 标题由文档添加，图表内不添加标题以免重复或样式不统一
            # ax.set_title(question_data.get('question_title', ''), fontsize=12, pad=20)
            
            plt.tight_layout()
            
            # 保存到临时文件
            temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            plt.savefig(temp_file.name, dpi=150, bbox_inches='tight')
            plt.close('all') # 关闭所有图表，释放内存
            
            return temp_file.name
            
        except Exception as e:
            print(f"Generate chart error: {e}")
            return None

    def _extract_dimensions_from_nodes(self, nodes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """从节点结构中提取维度数据"""
        dimensions = []
        
        def traverse(node_list):
            for node in node_list:
                # 检查该节点是否包含有效的数据和配置（作为维度）
                # 只要有数据且有图表配置或显示配置，就视为需要展示的“维度/问题”
                # 注意：这里我们放宽了限制，不再只看 data_analysis 类型
                is_valid_analysis = False
                
                # 如果是 data_analysis，或者有数据且有图表配置
                if node['type'] == 'data_analysis':
                    is_valid_analysis = True
                elif node.get('data') and (node.get('charts') or node.get('show_data_table')):
                    is_valid_analysis = True
                
                # 如果是容器或AI章节，且包含子节点，且子节点中有需要分析的内容
                if node.get('children'):
                    # 先看看自己是不是一个直接的分析节点（如果AI章节关联了图表）
                    if is_valid_analysis:
                        # 把它自己作为一个“维度”，其下的内容作为“问题”
                        # 但通常 AI 章节的图表是附属的。
                        # 这里为了简化，如果节点本身关联了数据，我们把它包装成一个维度+问题的结构
                        
                        # 构造一个单问题的维度
                        dim_item = {
                            "dimension_number": "",
                            "dimension_title": node['title'], # 维度名
                            "questions": [{
                                "question_number": node.get('data', {}).get('question_number') or 0,
                                "analysis_title": node['title'], # 问题名
                                "charts": node.get('charts', []),
                                "show_data_table": node.get('show_data_table', True),
                                "show_ai_interpretation": node.get('show_ai_interpretation', False),
                                "ai_analysis": node.get('content', '')
                            }]
                        }
                        dimensions.append(dim_item)
                    
                    # 然后继续递归查找子节点中的分析内容
                    # 注意：如果父节点已经是分析节点，子节点是否还需要提取？
                    # 假设结构是扁平的分析，或者嵌套的。
                    # 这里我们继续遍历子节点，可能会产生独立的维度。
                    # 为了避免重复（如果 traverse 逻辑是“提取所有分析”），我们需要小心。
                    # 现有的逻辑是：如果容器包含 data_analysis 子节点，则容器是维度。
                    
                    # 修正逻辑：
                    # 1. 检查是否有子节点是分析节点
                    analysis_children = []
                    for child in node.get('children', []):
                        if child['type'] == 'data_analysis' or (child.get('data') and (child.get('charts') or child.get('show_data_table'))):
                            analysis_children.append(child)
                    
                    if analysis_children:
                        # 如果有子分析节点，则当前节点作为维度容器
                        new_dim = {
                            "dimension_number": "", 
                            "dimension_title": node['title'],
                            "questions": []
                        }
                        for child in analysis_children:
                            q_data = child.get('data')
                            if not q_data:
                                continue
                            q_item = {
                                "question_number": q_data.get('question_number') or 0,
                                "analysis_title": child['title'],
                                "charts": child.get('charts', []),
                                "show_data_table": child.get('show_data_table', True),
                                "show_ai_interpretation": child.get('show_ai_interpretation', False),
                                "ai_analysis": child.get('content', '')
                            }
                            new_dim['questions'].append(q_item)
                        dimensions.append(new_dim)
                    
                    # 继续递归，查找深层嵌套（非直接子节点）
                    # 排除掉已经处理过的直接子节点
                    other_children = [c for c in node.get('children', []) if c not in analysis_children]
                    traverse(other_children)
                
                elif is_valid_analysis:
                    # 如果是叶子节点且是分析节点（且未被父级作为维度处理过，这在递归中很难判断）
                    # 在上面的逻辑中，父节点会处理它的直接分析子节点。
                    # 所以这里只需要处理那些没有被父节点捕获的情况吗？
                    # 其实上面的逻辑已经覆盖了“作为子节点”的情况。
                    # 这里唯一需要处理的是：如果顶层直接是一个 analysis 节点（没有父容器）
                    # 但 traverse 是从根列表开始的。
                    # 如果根列表里直接有一个 data_analysis，它应该被提取为一个独立的维度。
                    
                    # 构造独立维度
                    dim_item = {
                            "dimension_number": "",
                            "dimension_title": node['title'],
                            "questions": [{
                                "question_number": node.get('data', {}).get('question_number') or 0,
                                "analysis_title": node['title'],
                                "charts": node.get('charts', []),
                                "show_data_table": node.get('show_data_table', True),
                                "show_ai_interpretation": node.get('show_ai_interpretation', False),
                                "ai_analysis": node.get('content', '')
                            }]
                        }
                    dimensions.append(dim_item)

        traverse(nodes)
        return dimensions

    def _add_dimension_analysis(self, doc: Document, report_data: Dict[str, Any]):
        """添加维度分析"""
        heading = doc.add_heading('四、问卷结果分析', level=1)
        self._set_heading_style(heading, Pt(16))
        
        # 获取维度数据
        dimensions = report_data.get('dimensions', [])
        
        # 兼容新结构：如果 dimensions 为空，尝试从 generated_content.nodes 提取
        if not dimensions and 'generated_content' in report_data:
            nodes = report_data['generated_content'].get('nodes', [])
            if nodes:
                dimensions = self._extract_dimensions_from_nodes(nodes)

        excel_data = report_data.get('excel_data', {})
        questions_data = excel_data.get('questions', [])
        
        # 创建问题编号到数据的映射
        question_map = {q['question_number']: q for q in questions_data}
        
        for dim in dimensions:
            # 维度大标题
            dim_heading = doc.add_heading(
                f"{dim.get('dimension_number', '')} {dim['dimension_title']}", 
                level=2
            )
            self._set_heading_style(dim_heading, Pt(14))
            
            # 处理每个问题
            for q_config in dim.get('questions', []):
                q_number = q_config['question_number']
                
                # 小分析标题
                q_heading = doc.add_heading(
                    q_config['analysis_title'], 
                    level=3
                )
                self._set_heading_style(q_heading, Pt(12))
                
                # 获取对应的问题数据
                q_data = question_map.get(q_number)
                if not q_data:
                    continue
                
                # 获取内容展示顺序
                # content_order = q_config.get('content_order', ['table'])
                # if not content_order:
                #     content_order = ['table']
                # elif isinstance(content_order, str):
                #     content_order = [content_order]
                
                # 固定展示顺序：图表 -> 表格 -> AI解读
                content_order = ['chart', 'table', 'interpretation']

                # 按顺序插入内容
                for content_type in content_order:
                    if content_type == 'chart':
                        # 生成并插入图表
                        charts = q_config.get('charts', [])
                        if charts:
                            for chart_item in sorted(charts, key=lambda x: x.get('sort_order', 0)):
                                chart_path = self._generate_chart(q_data, chart_item)
                                if chart_path:
                                    try:
                                        doc.add_picture(chart_path, width=Inches(5.5))
                                        doc.add_paragraph()  # 空行
                                    except Exception as e:
                                        print(f"Error adding picture: {e}")
                                    finally:
                                        # 清理临时文件
                                        import os
                                        try:
                                            os.unlink(chart_path)
                                        except:
                                            pass
                                            
                    elif content_type == 'table':
                        # 显示数据表格
                        if q_config.get('show_data_table', True):
                            self._add_question_table(doc, q_data)
                            
                    elif content_type == 'interpretation':
                        # AI解读
                        if q_config.get('show_ai_interpretation', False):
                            ai_content = q_config.get('ai_analysis')
                            if ai_content:
                                p = doc.add_paragraph(ai_content)
                                self._set_paragraph_style(p)
    
    def _add_question_table(self, doc: Document, question_data: Dict[str, Any]):
        """添加问题数据表格"""
        options = question_data.get('options', [])
        
        if not options:
            return
        
        # 创建表格
        table = doc.add_table(rows=3, cols=len(options) + 2)
        table.style = 'Table Grid'
        
        # 表头行
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = question_data.get('question_title', '')
        hdr_cells[0].merge(hdr_cells[1])  # 合并前两列
        hdr_cells[0].text = question_data.get('question_title', '')
        
        for i, opt in enumerate(options):
            hdr_cells[i + 2].text = f"{opt['key']}. {opt['text']}"
        
        # 样本量行
        count_cells = table.rows[1].cells
        count_cells[0].text = ''
        count_cells[1].text = '样本量'
        for i, opt in enumerate(options):
            count_cells[i + 2].text = str(opt.get('count', 0))
        
        # 占比行
        pct_cells = table.rows[2].cells
        pct_cells[0].text = ''
        pct_cells[1].text = '占比'
        for i, opt in enumerate(options):
            pct_cells[i + 2].text = opt.get('percentage', '0%')
        
        # 空行
        doc.add_paragraph()
    
    def _add_attachments(self, doc: Document, report_data: Dict[str, Any]):
        """添加附件"""
        # 附件1：问卷原文
        heading = doc.add_heading('附件1：问卷原文', level=1)
        self._set_heading_style(heading, Pt(16))
        
        excel_data = report_data.get('excel_data', {})
        questions = excel_data.get('questions', [])
        
        for q in questions:
            p = doc.add_paragraph()
            run = p.add_run(f"{q['question_number']}. {q['question_title']}")
            run.bold = True
            
            for opt in q.get('options', []):
                doc.add_paragraph(f"   {opt['key']}. {opt['text']}", style='List Bullet')
        
        # 附件2：免责申明
        heading = doc.add_heading('附件2：免责申明', level=1)
        self._set_heading_style(heading, Pt(16))
        
        disclaimer = """（1）本次调研项目以随机选取对象进行面对面调研，本次调研只对本次样本数据负责。

（2）承接单位调研项目，是针对该产品品种调研，并非指定厂家指定品种。

（3）本次调研只针对调研区域数据负责，不代表全国调研数据。

服务提供单位：北京玖麟空科技有限公司"""
        
        for paragraph_text in disclaimer.split('\n\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                self._set_paragraph_style(p)
    
    def _set_heading_style(self, heading, font_size: Pt):
        """设置标题样式"""
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = font_size
            run.font.bold = True
    
    def _set_paragraph_style(self, paragraph):
        """设置段落样式"""
        for run in paragraph.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(12)
