#!/usr/bin/env python3
"""
豆包AI文档大纲识别测试脚本
"""
import asyncio
import sys
sys.path.insert(0, '/Users/lee/Documents/beiyi/爱诺模板项目/backend')

from docx import Document
from io import BytesIO
from app.services.doubao_service import DoubaoService


async def test_doubao_api():
    """测试豆包API是否正常工作"""
    print("=" * 60)
    print("测试豆包API连接")
    print("=" * 60)
    
    service = DoubaoService()
    
    try:
        # 简单测试
        response = await service.chat("你好，请用一句话介绍自己")
        print(f"✅ API连接成功")
        print(f"响应: {response[:100]}...")
        return True
    except Exception as e:
        print(f"❌ API连接失败: {e}")
        return False


async def test_structure_recognition():
    """测试文档大纲识别功能"""
    print("\n" + "=" * 60)
    print("测试文档大纲识别")
    print("=" * 60)
    
    # 创建一个测试文档
    doc = Document()
    
    # 添加各种格式的标题
    doc.add_heading('第一章 项目概述', level=1)
    doc.add_paragraph('这是正文内容，应该被忽略')
    
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph('这是正文内容，应该被忽略')
    
    doc.add_heading('1.1.1 研究目的', level=3)
    doc.add_paragraph('这是正文内容，应该被忽略')
    
    doc.add_heading('1.2 研究方法', level=2)
    doc.add_paragraph('这是正文内容，应该被忽略')
    
    doc.add_heading('第二章 市场分析', level=1)
    doc.add_heading('2.1 行业现状', level=2)
    doc.add_heading('2.1.1 市场规模', level=3)
    doc.add_heading('2.1.2 竞争格局', level=3)
    doc.add_heading('2.2 发展趋势', level=2)
    
    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # 使用AI提取器
    from app.services.ai_word_extractor import AIWordStructureExtractor
    
    extractor = AIWordStructureExtractor()
    
    try:
        result = await extractor.extract(buffer.getvalue(), "test_document.docx")
        
        print(f"✅ 文档解析成功")
        print(f"\n统计信息:")
        print(f"  - 总段落数: {result['stats']['total_paragraphs']}")
        print(f"  - 识别标题数: {result['stats']['recognized_headings']}")
        print(f"  - 一级标题: {result['stats']['by_level'][1]}")
        print(f"  - 二级标题: {result['stats']['by_level'][2]}")
        print(f"  - 三级标题: {result['stats']['by_level'][3]}")
        
        print(f"\n识别结构:")
        def print_structure(nodes, indent=0):
            for node in nodes:
                numbering = node.get('numbering', '')
                title = node.get('title', '')
                level = node.get('level', 0)
                prefix = "  " * indent
                num_str = f"[{numbering}] " if numbering else ""
                print(f"{prefix}{level}级: {num_str}{title}")
                
                children = node.get('children', [])
                if children:
                    print_structure(children, indent + 1)
        
        print_structure(result['structure'])
        
        return True
        
    except Exception as e:
        print(f"❌ 文档解析失败: {e}")
        import traceback
        traceback.print_exc()
        return False


async def main():
    """主函数"""
    print("\n" + "=" * 60)
    print("豆包AI文档大纲识别测试")
    print("=" * 60 + "\n")
    
    # 测试API连接
    api_ok = await test_doubao_api()
    
    if not api_ok:
        print("\n⚠️ API连接失败，无法进行文档识别测试")
        return
    
    # 测试文档识别
    recognition_ok = await test_structure_recognition()
    
    print("\n" + "=" * 60)
    if recognition_ok:
        print("✅ 所有测试通过！")
    else:
        print("❌ 测试失败，请检查配置")
    print("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())